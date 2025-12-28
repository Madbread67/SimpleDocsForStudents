from docx import Document
from datetime import datetime
import os
import re
from typing import Dict, List, Set

class UniversalDocumentGenerator:
    def __init__(self, template_dir='templates', output_dir='output'):
        self.template_dir = template_dir
        self.output_dir = output_dir

        # Создание директорий если их нет
        os.makedirs(template_dir, exist_ok=True)
        os.makedirs(output_dir, exist_ok=True)

        # Определение длины подчеркиваний для разных типов переменных
        self.placeholder_lengths = {
            'очень_короткое': 10,   # Например: курс, группа
            'короткое': 15,         # Например: день, месяц
            'среднее': 25,          # Например: телефон, год
            'длинное': 40,          # Например: ФИО, должность
            'очень_длинное': 60,    # Например: название организации
            'многострочное': 80     # Например: адрес
        }

    def _get_placeholder_length(self, variable_name):
        """Определяет длину подчеркивания на основе имени переменной"""
        variable_lower = variable_name.lower()

        # Очень короткие поля
        if any(word in variable_lower for word in ['курс', 'группа', 'день', 'месяц']):
            return self.placeholder_lengths['очень_короткое']

        # Короткие поля
        elif any(word in variable_lower for word in ['год', 'час', 'номер', 'код']):
            return self.placeholder_lengths['короткое']

        # Средние поля
        elif any(word in variable_lower for word in ['телефон', 'индекс', 'счет', 'оценка']):
            return self.placeholder_lengths['среднее']

        # Длинные поля
        elif any(word in variable_lower for word in ['фио', 'имя', 'фамилия', 'отчество']):
            return self.placeholder_lengths['длинное']

        # Очень длинные поля
        elif any(word in variable_lower for word in ['название', 'организация', 'специальность', 'модуль']):
            return self.placeholder_lengths['очень_длинное']

        # Многострочные
        elif any(word in variable_lower for word in ['адрес', 'описание', 'задание', 'требование']):
            return self.placeholder_lengths['многострочное']

        # По умолчанию
        return self.placeholder_lengths['среднее']

    def scan_template_for_variables(self, template_path: str) -> Set[str]:
        """Сканирует шаблон и возвращает множество найденных переменных"""
        variables = set()

        try:
            doc = Document(template_path)

            # Регулярное выражение для поиска переменных в формате {{variable}}
            variable_pattern = r'\{\{([^}]+)\}\}'

            # Поиск в параграфах
            for paragraph in doc.paragraphs:
                matches = re.findall(variable_pattern, paragraph.text)
                variables.update(matches)

            # Поиск в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = re.findall(variable_pattern, cell.text)
                        variables.update(matches)

                        # Также ищем в параграфах внутри ячеек
                        for paragraph in cell.paragraphs:
                            matches = re.findall(variable_pattern, paragraph.text)
                            variables.update(matches)

            # Поиск в верхних/нижних колонтитулах
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    matches = re.findall(variable_pattern, paragraph.text)
                    variables.update(matches)
                for paragraph in section.footer.paragraphs:
                    matches = re.findall(variable_pattern, paragraph.text)
                    variables.update(matches)

        except Exception as e:
            print(f"Ошибка при сканировании шаблона {template_path}: {e}")

        return variables

    def extract_all_variables_from_templates(self) -> Dict[str, Set[str]]:
        """Извлекает все переменные из всех шаблонов"""
        templates_vars = {}

        if not os.path.exists(self.template_dir):
            return templates_vars

        for template_file in os.listdir(self.template_dir):
            if template_file.endswith('.docx'):
                template_path = os.path.join(self.template_dir, template_file)
                variables = self.scan_template_for_variables(template_path)
                if variables:
                    templates_vars[template_file] = variables

        return templates_vars

    def prepare_student_data(self, student_data: Dict) -> Dict[str, str]:
        """Подготовка данных студента для всех возможных переменных"""
        # Базовые данные из БД
        data = {
            # Основные данные студента
            'FIO': student_data.get('full_name', ''),
            'ФИО': student_data.get('full_name', ''),
            'Specialnost_id': student_data.get('specialty_id', ''),
            'Specialnost': student_data.get('specialty_name', ''),
            'Grupa': student_data.get('group_name', ''),
            'Модуль': student_data.get('module_name', ''),

            # Даты практики
            'дата начала практики': f"{student_data.get('practice_start_day', '')} {student_data.get('practice_start_month', '')} {student_data.get('practice_start_year', '')}",
            'дата конца практики': f"{student_data.get('practice_end_day', '')} {student_data.get('practice_end_month', '')} {student_data.get('practice_end_year', '')}",
            'practice_start_day': student_data.get('practice_start_day', ''),
            'practice_start_month': student_data.get('practice_start_month', ''),
            'practice_start_year': student_data.get('practice_start_year', ''),
            'practice_end_day': student_data.get('practice_end_day', ''),
            'practice_end_month': student_data.get('practice_end_month', ''),
            'practice_end_year': student_data.get('practice_end_year', ''),

            # Организация
            'Организация Практики': student_data.get('organization_name', ''),
            'Адрес Организации': student_data.get('organization_address', ''),
            'Адрес Практики': student_data.get('organization_address', ''),

            # Преподаватель
            'Преподаватель_ФИО': student_data.get('teacher_name', ''),
            'Преподаватель_телефон': student_data.get('teacher_phone', ''),

            # Модуль и часы
            'Модуль_часы': student_data.get('module_hours', ''),
            'practice_hours': student_data.get('practice_hours', ''),

            # Дополнительные поля (могут быть вычислены)
            'Текущая_дата': datetime.now().strftime('%d.%m.%Y'),
            'Текущий_год': datetime.now().strftime('%Y'),
        }

        # Вычисляемые поля
        full_name = data['FIO']
        if full_name:
            parts = full_name.split()
            if len(parts) >= 3:
                data['Инициалы'] = f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
                data['Фамилия'] = parts[0]
                data['Имя'] = parts[1] if len(parts) > 1 else ''
                data['Отчество'] = parts[2] if len(parts) > 2 else ''
            elif len(parts) == 2:
                data['Инициалы'] = f"{parts[0]} {parts[1][0]}."
                data['Фамилия'] = parts[0]
                data['Имя'] = parts[1]
            elif len(parts) == 1:
                data['Инициалы'] = parts[0]
                data['Фамилия'] = parts[0]

        # Дни практики (вычисляем разницу между датами)
        try:
            start_year = int(student_data.get('practice_start_year', 2025))
            end_year = int(student_data.get('practice_end_year', 2025))

            # Простой расчет дней (10 рабочих дней как в шаблоне)
            data['Количество_дней'] = '10'
            data['Пропущено_дней'] = '0'
            data['По_неуважительной_причине'] = '0'
        except:
            data['Количество_дней'] = '10'
            data['Пропущено_дней'] = '0'
            data['По_неуважительной_причине'] = '0'

        # Для ваших шаблонов добавляем специфичные поля
        data['Kurs'] = '3'  # Можно вычислять из группы или добавлять в БД
        data['FormaObucheniaOchnaia_or_Zaochnaya'] = 'Очная'

        # Удаляем None значения и преобразуем в строки
        for key in list(data.keys()):
            if data[key] is None:
                data[key] = ''
            else:
                data[key] = str(data[key]).strip()

        return data

    def _replace_all_placeholders(self, doc, variables_data):
        """Замена всех заполнителей в документе"""
        # Регулярное выражение для поиска переменных
        variable_pattern = r'\{\{([^}]+)\}\}'

        # Замена в параграфах
        for paragraph in doc.paragraphs:
            if '{{' in paragraph.text:
                # Находим все переменные в тексте
                matches = re.findall(variable_pattern, paragraph.text)
                for match in matches:
                    # Получаем значение или подчеркивание
                    value = variables_data.get(match.strip())
                    if value is None or value == '':
                        length = self._get_placeholder_length(match)
                        value = '_' * length

                    # Заменяем
                    placeholder = f'{{{{{match}}}}}'
                    paragraph.text = paragraph.text.replace(placeholder, value)

                    # Также заменяем в отдельных runs для сохранения форматирования
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '{{' in cell.text:
                        # Находим все переменные в тексте ячейки
                        matches = re.findall(variable_pattern, cell.text)
                        for match in matches:
                            # Получаем значение или подчеркивание
                            value = variables_data.get(match.strip())
                            if value is None or value == '':
                                length = self._get_placeholder_length(match)
                                value = '_' * length

                            # Заменяем
                            placeholder = f'{{{{{match}}}}}'
                            cell.text = cell.text.replace(placeholder, value)

                            # Также заменяем в параграфах внутри ячеек
                            for paragraph in cell.paragraphs:
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, value)

        # Замена в верхних/нижних колонтитулах
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                if '{{' in paragraph.text:
                    matches = re.findall(variable_pattern, paragraph.text)
                    for match in matches:
                        value = variables_data.get(match.strip())
                        if value is None or value == '':
                            length = self._get_placeholder_length(match)
                            value = '_' * length

                        placeholder = f'{{{{{match}}}}}'
                        paragraph.text = paragraph.text.replace(placeholder, value)

            for paragraph in section.footer.paragraphs:
                if '{{' in paragraph.text:
                    matches = re.findall(variable_pattern, paragraph.text)
                    for match in matches:
                        value = variables_data.get(match.strip())
                        if value is None or value == '':
                            length = self._get_placeholder_length(match)
                            value = '_' * length

                        placeholder = f'{{{{{match}}}}}'
                        paragraph.text = paragraph.text.replace(placeholder, value)

    def generate_document(self, template_name: str, student_data: Dict, output_filename: str = None) -> str:
        """Генерация документа по шаблону с автоподстановкой всех переменных"""
        try:
            template_path = os.path.join(self.template_dir, template_name)
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"Шаблон не найден: {template_path}")

            # Загружаем документ
            doc = Document(template_path)

            # Сканируем шаблон для отладки (можно убрать в продакшене)
            variables_in_template = self.scan_template_for_variables(template_path)
            print(f"Найдены переменные в {template_name}: {variables_in_template}")

            # Подготавливаем данные студента
            prepared_data = self.prepare_student_data(student_data)

            # Заменяем все переменные
            self._replace_all_placeholders(doc, prepared_data)

            # Генерация имени файла если не указано
            if not output_filename:
                safe_name = re.sub(r'[<>:"/\\|?*]', '_', student_data.get('full_name', 'Неизвестно'))
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                base_name = os.path.splitext(template_name)[0]
                output_filename = f"{base_name}_{safe_name}_{timestamp}.docx"

            output_path = os.path.join(self.output_dir, output_filename)
            doc.save(output_path)

            print(f"✓ Документ сохранен: {output_path}")
            return output_path

        except Exception as e:
            print(f"✗ Ошибка генерации документа {template_name}: {e}")
            raise

    def generate_all_documents(self, student_data: Dict) -> List[str]:
        """Генерация всех доступных документов для студента"""
        if not os.path.exists(self.template_dir):
            raise FileNotFoundError(f"Папка с шаблонами не найдена: {self.template_dir}")

        results = []

        for template_file in os.listdir(self.template_dir):
            if template_file.endswith('.docx'):
                try:
                    print(f"Генерация документа: {template_file}")
                    output_path = self.generate_document(template_file, student_data)
                    results.append((template_file, output_path, 'success'))
                except Exception as e:
                    results.append((template_file, str(e), 'error'))
                    print(f"Ошибка при генерации {template_file}: {e}")

        return results

    def get_available_templates(self) -> List[str]:
        """Получение списка доступных шаблонов"""
        templates = []
        if os.path.exists(self.template_dir):
            for file in os.listdir(self.template_dir):
                if file.endswith('.docx'):
                    templates.append(file)
        return sorted(templates)

    def get_template_variables(self, template_name: str) -> Set[str]:
        """Получение списка переменных для конкретного шаблона"""
        template_path = os.path.join(self.template_dir, template_name)
        if os.path.exists(template_path):
            return self.scan_template_for_variables(template_path)
        return set()

    def get_all_variables_summary(self) -> Dict[str, Set[str]]:
        """Получение сводки по всем переменным во всех шаблонах"""
        return self.extract_all_variables_from_templates()
